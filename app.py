from flask import Flask, request, render_template, send_file
import docx
import csv
import os

app = Flask(__name__)

@app.route('/')
def home():
    return '''
    <!doctype html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Upload DOCX File</title>
    </head>
    <body>
        <h1>Upload a DOCX file to extract the first table</h1>
        <form action="/extract" method="post" enctype="multipart/form-data">
            <input type="file" name="docx_file" accept=".docx" required>
            <button type="submit">Upload and Extract</button>
        </form>
    </body>
    </html>
    '''

@app.route('/extract', methods=['POST'])
def extract_first_table():
    if 'docx_file' not in request.files:
        return "No file part", 400

    docx_file = request.files['docx_file']
    if docx_file.filename == '':
        return "No selected file", 400

    # Load the .docx document
    doc = docx.Document(docx_file)

    # Get the first table from the document
    if not doc.tables:
        return "No tables found in the document", 400

    table = doc.tables[0]

    # Extract data from the table
    data = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        data.append(row_data)

    # Save the data to a CSV file
    csv_file = 'extracted_table.csv'
    with open(csv_file, mode='w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerows(data)

    return send_file(csv_file, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)
